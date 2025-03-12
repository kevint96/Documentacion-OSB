import streamlit as st
import os
import shutil
from zipfile import ZipFile
import zipfile
import tempfile
import subprocess
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import inspect
import os
import xml.etree.ElementTree as ET
import gspread
import time  # Importar el módulo time
import logging
import re
import inspect
import ast
from datetime import datetime
import difflib
import glob

def print_with_line_number(msg):
    caller_frame = inspect.currentframe().f_back
    line_number = caller_frame.f_lineno
    print(f"Linea {line_number}: {msg}")
    print("")
    
def apply_format(run,fuente,size,negrita,color):
    run.font.name = fuente  # Cambiar el nombre de la fuente
    run.font.size = Pt(size)  # Cambiar el tamaño de la fuente
    run.font.bold = negrita  # Aplicar negrita
    run.font.color.rgb = RGBColor(0, 0, color)  # Cambiar el color del texto a rojo

def replace_text_in_paragraph(paragraph, replacements):
    full_text = paragraph.text
    contador = 1
    #st.success(f"Texto en linea: {full_text}")
    for key, value in replacements.items():
        if key in full_text:
            #st.success(f"full_text: {full_text}")
            #st.success(f"p paragraphs: {paragraph.text}")
            #st.success(f"clave coincide: {key}")
            full_text = full_text.replace(key, str(value))  # Actualiza full_text
            
            if key in '{nombre_servicio_inicial}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial',18,True,0)  # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
            if key in '{nombre_operacion_inicial}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial',10,True,0)  # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
            if key in '{nombre_servicio_secundario}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial',10,True,0)    # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            if key in '{nombre_operacion}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Times New Roman',10,False,0)    # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            if key in '{unique_operations}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Times New Roman',10,False,0)    # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            if key in '{nombre_servicio}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Times New Roman',10,False,0)    # Aplicar formato al texto del párrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            if key in '{nombre_servicio_contrato}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Times New Roman',10,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{nombre_servicio_wsdl}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Times New Roman',10,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{nombre_servicio_contrato2}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial',10,False,0)  # Aplicar formato al texto del párrafo
                
            if key in '{nombre_servicio_tabla}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',11,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{fecha}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{autor_inicial}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,True,0)  # Aplicar formato al texto del párrafo
            
            if key in '{autor}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial',10,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{autor2}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,False,0)  # Aplicar formato al texto del párrafo
            
            if key in '{url}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,False,255)  # Aplicar formato al texto del párrafo
                
            if key in '{operacion_legado}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,False,255)  # Aplicar formato al texto del párrafo
                
            
            if key in '{proyecto_abc}':
                paragraph.clear()  # Limpiar el párrafo
                paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                apply_format(paragraph.runs[0],'Arial MT',10,False,0)  # Aplicar formato al texto del párrafo

def print_element_content(element, element_name):
    st.success(f"Contenido del {element_name}:")
    for paragraph in element.paragraphs:
        print_with_line_number(paragraph.text)
    for table in element.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print_with_line_number(paragraph.text)

def replace_text_in_element(element, replacements):
    for paragraph in element.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    for table in element.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

def replace_text_in_doc(doc, replacements):
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    for section in doc.sections:
        st.success(f"Encabezado de la sección: {section.header}")
        print_element_content(section.header, "Encabezado de la sección")
        replace_text_in_element(section.header, replacements)
        st.success(f"Pie de página de la sección: {section.footer}")
        print_element_content(section.footer, "Pie de página de la sección")
        replace_text_in_element(section.footer, replacements)
        # Agregamos este bloque específico para procesar las tablas dentro del encabezado de la sección 2
        if "Encabezado-Sección 2-" in [paragraph.text for paragraph in section.header.paragraphs]:
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            print_with_line_number(paragraph.text)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, replacements)
    
    doc = reemplazar_texto_en_doc(doc, replacements)
    
    return doc

   
def service_refs_ruta_pipeline(pipeline_path, project_path):
    
    elemento = ""
    # Servicios a excluir
    servicios_a_excluir = [
        'ComponentesComunes/Proxies/PS_ManejadorGenericoErroresV1.0',
        'UtilitariosEBS/Proxies/AuditoriaSOA/RegistrarAuditoriaSOADATV1.0'
    ]
    
    while True:
        
        st.success(f"pipeline_path: {pipeline_path}")
        
        # Leer el archivo .pipeline
        with open(pipeline_path, 'r') as file:
            pipeline_content = file.readlines()

        # Buscar todas las líneas que contienen ':service ref="'
        matching_lines = [line for line in pipeline_content if ':service ref="' in line]

        # Extraer la información deseada de las líneas coincidentes
        servicios = set()  # Usamos un conjunto para evitar elementos duplicados
        for line in matching_lines:
            service_start_index = line.find(':service ref="') + len(':service ref="')
            service_end_index = line.find('"', service_start_index)
            service_ref = line[service_start_index:service_end_index]
            # Verificar si el servicio no está en la lista de servicios a excluir
            if service_ref not in servicios_a_excluir:
                servicios.add(service_ref)

        # Imprimir los servicios encontrados
        print_with_line_number("Servicios encontrados:")
        for service in servicios:
            print_with_line_number(service)
            
             # Si el elemento contiene '/BusinessServices/', salir del bucle
            if '/BusinessServices/' in service:
                st.success(f"BusinessServices: {service}")
                business_path = os.path.join(project_path, service + '.bix')
                
                with open(business_path, 'r') as business_file:
                    business_content = business_file.readlines()
                    
                    matching_lines = [line for line in business_content if 'operation-name>' in line]
                    
                    # Extraer los elementos ref de las líneas coincidentes
                    elementos_ref = set()  # Usamos un conjunto para evitar elementos duplicados
                    for line in matching_lines:
                        invoke_start_index = line.find('operation-name>') + len('operation-name>')
                        invoke_end_index = line.find('<', invoke_start_index)
                        invoke_ref = line[invoke_start_index:invoke_end_index]
                        elementos_ref.add(invoke_ref)

                    # Imprimir los elementos ref encontrados
                    print_with_line_number("Elementos ref encontrados en {}: ".format(service))
                    for elemento in elementos_ref:
                        print_with_line_number(elemento)
                return elemento

            # Construir la ruta del archivo proxy
            proxy_path = os.path.join(project_path, service + '.proxy')

            # Verificar si el archivo proxy existe
            if os.path.exists(proxy_path):
                # Leer el archivo proxy
                with open(proxy_path, 'r') as proxy_file:
                    proxy_content = proxy_file.readlines()

                # Buscar todas las líneas que contienen ':invoke ref="'
                matching_lines = [line for line in proxy_content if ':invoke ref="' in line]

                # Extraer los elementos ref de las líneas coincidentes
                elementos_ref = set()  # Usamos un conjunto para evitar elementos duplicados
                for line in matching_lines:
                    invoke_start_index = line.find(':invoke ref="') + len(':invoke ref="')
                    invoke_end_index = line.find('"', invoke_start_index)
                    invoke_ref = line[invoke_start_index:invoke_end_index]
                    elementos_ref.add(invoke_ref)

                # Imprimir los elementos ref encontrados
                print_with_line_number("Elementos ref encontrados en {}: ".format(service))
                for elemento in elementos_ref:
                    print_with_line_number(elemento)

                    # Si el elemento contiene '/BusinessServices/', salir del bucle
                    if '/BusinessServices/' in elemento:
                        st.success(f"elemento: {elemento}")
                        return elemento
                    else:
                        pipeline_path = os.path.join(project_path, elemento + '.pipeline')
                       
            else:
                print_with_line_number("El archivo proxy {} no existe.".format(proxy_path))
                break

    return elemento

def extract_xsd_import_paths(wsdl_path):
    xsd_import_paths = set()  # Conjunto para almacenar rutas únicas

    # Leer el contenido del archivo WSDL
    with open(wsdl_path, 'r', encoding='utf-8') as file:
        wsdl_content = file.read()

    # Extraer el contenido dentro de CDATA usando una expresión regular
    cdata_match = re.search(r'<!\[CDATA\[(.*?)\]\]>', wsdl_content, re.DOTALL)
    
    if cdata_match:
        cdata_content = cdata_match.group(1)  # Obtener solo el contenido dentro de CDATA

        # Parsear el contenido XML dentro del CDATA
        try:
            root = ET.fromstring(cdata_content)  # Convertir el CDATA en XML
        except ET.ParseError as e:
            print("Error al parsear el contenido de CDATA:", e)
            return xsd_import_paths

        # Espacios de nombres comunes en WSDL
        namespaces = {
            'xsd': 'http://www.w3.org/2001/XMLSchema'
        }

        # Buscar todos los elementos <xsd:import>
        for xsd_import in root.findall(".//xsd:import", namespaces):
            schema_location = xsd_import.get("schemaLocation")
            if schema_location:
                xsd_import_paths.add(schema_location)
    return list(xsd_import_paths)  # Convertimos el conjunto de vuelta a lista antes de devolverlo

def find_import_elements_with_namespace(xsd_content, target_namespace, xsd_file_path):
    schema_location = ""
    absolute_schema_location = None  # Inicializa la variable

    namespaces = {
        'xsd': 'http://www.w3.org/2001/XMLSchema'
        # Agrega otros namespaces si es necesario
    }
    st.success(f"target_namespace: {target_namespace}")

    root = ET.fromstring(xsd_content)
    
    st.success(f"xsd_file_path: {xsd_file_path}")
    
    # Busca todos los elementos import
    xsd_import_elements = root.findall(".//xsd:import", namespaces)

    for import_element in xsd_import_elements:
        namespace = import_element.get('namespace')
        st.success(f"namespace: {namespace}")
        if namespace == target_namespace:
            schema_location = import_element.get('schemaLocation')
            st.success(f"Found xsd:import with namespace '{namespace}': {schema_location}")
            
            # Concatena la ruta del archivo XSD principal con la ubicación del esquema importado
            absolute_schema_location = os.path.normpath(os.path.join(os.path.dirname(xsd_file_path), schema_location)).replace('\\', '/')
            st.success(f"schema_location: {absolute_schema_location}")
            break  # Si encuentras la coincidencia, sal del bucle
    
    return absolute_schema_location  # Esto devolverá None si no se encontró coincidencia "

def extract_namespaces(xsd_content):
    """Extrae los namespaces definidos en el XSD."""
    namespaces = {}
    matches = re.findall(r'xmlns:([\w]+)="([^"]+)"', xsd_content)
    for prefix, uri in matches:
        namespaces[prefix] = uri
    return namespaces


def extract_imports(root):
    """Extrae los imports y los mapea con sus schemaLocation."""
    # Detectar el prefijo correcto para XML Schema (puede ser xs: o xsd:)
    schema_ns = "http://www.w3.org/2001/XMLSchema"
    prefix = None
    
    # Buscar en los atributos del root el namespace correspondiente
    for attr in root.attrib:
        if root.attrib[attr] == schema_ns:
            prefix = attr.split(":")[-1]  # Extraer el prefijo después de "xmlns:"
            break
    
    # Si no se encontró prefijo, usar xs por defecto
    if not prefix:
        prefix = "xs"

    # Buscar los imports con el prefijo detectado
    imports = {}
    for imp in root.findall(f".//{prefix}:import", {prefix: schema_ns}):
        namespace = imp.attrib.get('namespace')
        schema_location = imp.attrib.get('schemaLocation')
        if namespace and schema_location:
            imports[namespace] = schema_location
    
    return imports

def get_correct_xsd_path(current_xsd_path, schema_location):
    """
    Corrige la ruta de un XSD importado considerando los niveles de directorio.
    """
    base_path = os.path.dirname(current_xsd_path)  # Obtener la carpeta del XSD actual
    st.success(f"base_path: {base_path}")
    corrected_path = os.path.abspath(os.path.join(base_path, schema_location))
    st.success(f"corrected_path: {corrected_path}")    # Resolver la ruta correcta
    corrected_path = corrected_path.replace("/mount/src/documentacion-osb/extraccion_jar","")
    corrected_path = corrected_path.replace(".xsd",".XMLSchema")
    st.success(f"corrected_path: {corrected_path}")    # Resolver la ruta correcta

    return corrected_path

def parse_xsd_file(project_path, xsd_file_path, operation_name, service_url, capa_proyecto, 
                   operacion_business, operations, service_name, operation_actual, 
                   target_complex_type=None, root_element_name=None):
    """
    Parsea un XSD y extrae los elementos request/response de forma recursiva.
    """
    request_elements = []
    response_elements = []

    extraccion_dir = os.path.abspath(project_path)
    xsd_file_path = os.path.normpath(xsd_file_path.strip("/\\"))  
    subcarpeta_xsd = os.path.dirname(xsd_file_path)
    subcarpeta_xsd = os.path.normpath(subcarpeta_xsd).replace("../", "")

    ruta_corregida = os.path.join(extraccion_dir, subcarpeta_xsd, os.path.basename(xsd_file_path))
    
    st.success(f"Ruta corregida FINAL: {ruta_corregida}")
    
    if not os.path.isfile(ruta_corregida):
        st.error(f"El archivo XSD {ruta_corregida} no existe.")
        return request_elements, response_elements

    # Leer el contenido del XSD
    with open(ruta_corregida, 'r', encoding="utf-8") as f:
        xsd_content = f.read()

    # Extraer el contenido de CDATA si es necesario
    cdata_match = re.search(r'<!\[CDATA\[(.*?)\]\]>', xsd_content, re.DOTALL)
    if cdata_match:
        xsd_content = cdata_match.group(1)
        st.success("Se ha extraído el contenido de CDATA correctamente")

    try:
        root = ET.fromstring(xsd_content)
    except ET.ParseError as e:
        st.error(f"Error al analizar el XMLSchema: {e}")
        return request_elements, response_elements

    namespaces = extract_namespaces(xsd_content)
    imports = extract_imports(root)

    st.success(f"Namespaces detectados: {namespaces}")
    st.success(f"Imports encontrados: {imports}")

    # Obtener todos los complexTypes
    complex_types = {elem.attrib.get('name', None): elem for elem in root.findall(".//xs:complexType", namespaces) if 'name' in elem.attrib}

    # ✅ Buscar los elementos principales del XSD (Request y Response)
    root_elements = {elem.attrib.get('name', ''): elem.attrib.get('type', '').split(':')[-1] for elem in root.findall(".//xs:element", namespaces)}

    # 🚀 **Si `target_complex_type` está definido, buscar SOLO ese complexType.**
    if target_complex_type:
        st.success(f"🔍 Buscando SOLO el complexType: {target_complex_type}")
        explorar_complex_type(target_complex_type, root_element_name, complex_types, namespaces, imports, extraccion_dir, 
                              xsd_file_path, project_path, service_url, capa_proyecto, operacion_business, 
                              operations, service_name, operation_actual, request_elements, response_elements,operation_name)
        return request_elements, response_elements

    # 🔹 Si `target_complex_type` no está, procesamos TODO desde los elementos raíz.
    for root_element_name, complex_type in root_elements.items():
        st.success(f"Procesando raíz: {root_element_name} -> {complex_type}")

        if complex_type in complex_types:
            explorar_complex_type(complex_type, root_element_name, complex_types, namespaces, imports, extraccion_dir, 
                                  xsd_file_path, project_path, service_url, capa_proyecto, operacion_business, 
                                  operations, service_name, operation_actual, request_elements, response_elements,operation_name)

    return request_elements, response_elements


def explorar_complex_type(type_name, parent_element_name, complex_types, namespaces, imports, extraccion_dir, 
                          xsd_file_path, project_path, service_url, capa_proyecto, operacion_business, 
                          operations, service_name, operation_actual, request_elements, response_elements,operation_name):
    """Explora recursivamente un complexType y extrae sus elementos internos."""

    type_name = type_name.split(':')[-1]  

    if type_name in complex_types:
        st.success(f"Explorando complexType: {type_name}")
        sequence = complex_types[type_name].find('xs:sequence', namespaces)

        if sequence is not None:
            for element in sequence.findall('xs:element', namespaces):
                element_name = element.attrib.get('name', '')
                element_type = element.attrib.get('type', '')

                # 🔹 **Asegurar que `parent_element_name` se mantenga para preservar la jerarquía**
                full_name = f"{parent_element_name}.{element_name}" if parent_element_name else element_name
                st.success(f"Encontrado elemento: {full_name} con tipo: {element_type}")

                simple_type = element.find('xs:simpleType', namespaces)
                if simple_type is not None:
                    restriction = simple_type.find('xs:restriction', namespaces)
                    if restriction is not None and 'base' in restriction.attrib:
                        element_type = restriction.attrib['base']
                        st.success(f"Elemento {full_name} tiene restricción con base: {element_type}")

                if element_type.startswith(("xsd:", "xs:")):
                    element_details = {
                        'elemento': parent_element_name,
                        'name': full_name,
                        'type': element_type,
                        'url': service_url,
                        'ruta': capa_proyecto,
                        'business': operacion_business,
                        'operations': operations,
                        'service_name': service_name,
                        'operation_actual': operation_actual,
                    }
                    st.success(f"Agregando elemento primitivo: {element_details}")
                    if 'Request' in parent_element_name:
                        request_elements.append(element_details)
                    elif 'Response' in parent_element_name:
                        response_elements.append(element_details)
                        
                # ✅ Si es otro complexType dentro del mismo XSD
                elif element_type in complex_types:
                    st.success(f"Buscando {element_type} en el mismo XSD")
                    explorar_complex_type(element_type, full_name, complex_types, namespaces, imports, extraccion_dir, 
                                          xsd_file_path, project_path, service_url, capa_proyecto, operacion_business, 
                                          operations, service_name, operation_actual, request_elements, response_elements,operation_name)

                elif ':' in element_type:
                    prefix, nested_type = element_type.split(':')
                    
                    if nested_type in complex_types:
                        st.success(f"Buscando {nested_type} en el mismo XSD")
                        explorar_complex_type(nested_type, full_name, complex_types, namespaces, imports, extraccion_dir, 
                                              xsd_file_path, project_path, service_url, capa_proyecto, operacion_business, 
                                              operations, service_name, operation_actual, request_elements, response_elements,operation_name)
                    elif prefix in namespaces:
                        namespace = namespaces[prefix]
                        if namespace in imports:
                            schema_location = imports[namespace]
                            st.warning(f"El tipo {nested_type} está en otro XSD: {schema_location}")
                            corrected_xsd_path = get_correct_xsd_path(xsd_file_path, schema_location)
                            new_xsd_path = os.path.join(extraccion_dir, corrected_xsd_path)

                            # ✅ Se envía `target_complex_type` en la recursión
                            parse_xsd_file(project_path, new_xsd_path, operation_name, service_url, 
                                           capa_proyecto, operacion_business, operations, 
                                           service_name, operation_actual, 
                                           target_complex_type=nested_type, 
                                           root_element_name=full_name)
                    else:
                        st.warning(f"No se encontró el namespace para el prefijo {prefix}")
                else:
                    st.warning(f"complexType {element_type} no encontrado en el XSD")
    else:
        st.warning(f"complexType {type_name} no encontrado en el XSD")

def leer_xsd_file(xsd_file_path, complexType_name):
    elements_list = []

    if xsd_file_path.endswith('.xsd') and os.path.isfile(xsd_file_path):
        with open(xsd_file_path, 'r', encoding="utf-8") as f:
            xsd_content = f.read()
            root = ET.fromstring(xsd_content)
            namespaces = {'xs': 'http://www.w3.org/2001/XMLSchema'}
            
            st.success(f"xsd_file_path: {xsd_file_path}")
            print_with_line_number("")

            # Función para detectar y eliminar repeticiones cíclicas en los nombres de los elementos
            def remove_repetitions(element_name):
                parts = element_name.split('.')
                seen = set()
                unique_parts = []
                for part in parts:
                    if part in seen:
                        break
                    seen.add(part)
                    unique_parts.append(part)
                return '.'.join(unique_parts)

            # Función para obtener elementos recursivamente con control de visitas
            def get_elements(complex_type_element, parent_name, visited):
                sequence_element = complex_type_element.find('xs:sequence', namespaces)
                if sequence_element is not None:
                    child_elements = sequence_element.findall('xs:element', namespaces)
                    for child_element in child_elements:
                        element_name = child_element.attrib.get('name', '')
                        element_type = child_element.attrib.get('type', '')
                        full_element_name = f"{parent_name}.{element_name}"

                        # Detectar y eliminar repeticiones cíclicas
                        full_element_name = remove_repetitions(full_element_name)

                        st.success(f"element_name: {full_element_name}")
                        st.success(f"element_type: {element_type}")
                        if not element_type:
                            element_type = 'xs:string'
                        elements_list.append({'element_name': full_element_name, 'element_type': element_type})

                        if ':' in element_type:
                            prefix, complexType_name_interno = element_type.split(':')
                            if complexType_name_interno not in visited:
                                visited.add(complexType_name_interno)
                                complex_type_element = root.find(f".//xs:complexType[@name='{complexType_name_interno}']", namespaces)
                                if complex_type_element is not None:
                                    get_elements(complex_type_element, full_element_name, visited)

            complex_type_element = root.find(f".//xs:complexType[@name='{complexType_name}']", namespaces)
            if complex_type_element is not None:
                print_with_line_number("")
                st.success(f"complex_type_name: {complexType_name}")
                print_with_line_number("")
                st.success(f"complex_type_element: {complex_type_element}")
                print_with_line_number("")
                
                visited = set()
                get_elements(complex_type_element, complexType_name, visited)
                
    return elements_list
    
def has_http_provider_id(xml_content):
    root = ET.fromstring(xml_content)
    namespaces = {'tran': 'http://www.bea.com/wli/sb/transports'}
    provider_id_element = root.find(".//tran:provider-id", namespaces)
    return provider_id_element is not None and provider_id_element.text == 'http'

def extract_project_name_from_proxy(proxy_path):
    try:
        with open(proxy_path, 'r', encoding="utf-8") as f:
            content = f.read()
            start = content.find('<con:wsdl ref="') + len('<con:wsdl ref="')
            end = content.find('"', start)
            wsdl_ref = content[start:end]
            return wsdl_ref.split("/")[0]
    except FileNotFoundError:
        #st.success(f"El archivo {proxy_path} no existe.")
        return None

def reemplazar_texto_en_doc(doc, reemplazos):
    """
    Reemplaza variables en el documento, incluyendo encabezados, pies de página y contenido.
    """
    # Reemplazo en párrafos normales
    for parrafo in doc.paragraphs:
        for clave, valor in reemplazos.items():
            if clave in parrafo.text:
                parrafo.text = parrafo.text.replace(clave, valor)
    
    # Reemplazo en encabezados y pies de página
    for section in doc.sections:
        # Encabezado
        for parrafo in section.header.paragraphs:
            for clave, valor in reemplazos.items():
                if clave in parrafo.text:
                    parrafo.text = parrafo.text.replace(clave, valor)
        
        # Pie de página
        for parrafo in section.footer.paragraphs:
            for clave, valor in reemplazos.items():
                if clave in parrafo.text:
                    parrafo.text = parrafo.text.replace(clave, valor)
    
    # Reemplazo en tablas sin alterar el formato
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for clave, valor in reemplazos.items():
                    if clave in celda.text:
                        celda.text = celda.text.replace(clave, valor)
    
    return doc

def extract_service_url(xml_content):
    root = ET.fromstring(xml_content)
    tran_namespace = {'tran': 'http://www.bea.com/wli/sb/transports', 'env': 'http://www.bea.com/wli/config/env'}
    uri_element = root.find(".//tran:URI/env:value", namespaces=tran_namespace)
    if uri_element is not None:
        return uri_element.text
    return ''

def extract_pipeline_path_from_proxy(proxy_path, jdeveloper_projects_dir):
    try:
        with open(proxy_path, 'r', encoding="utf-8") as f:
            content = f.read()
            start = content.find('<ser:invoke ref="') + len('<ser:invoke ref="')
            end = content.find('"', start)
            pipeline_ref = content[start:end]
            pipeline_path = os.path.join(jdeveloper_projects_dir, pipeline_ref + ".Pipeline")
            return pipeline_path
    except FileNotFoundError:
        print(f"El archivo {proxy_path} no pudo ser encontrado.")
        return None  # O puedes lanzar otra excepción, dependiendo del flujo de tu programa.
     
def extract_wsdl_relative_path(xml_content):
    root = ET.fromstring(xml_content)
    namespaces = {'con': 'http://www.bea.com/wli/sb/services/bindings/config'}
    wsdl_ref_element = root.find(".//con:wsdl", namespaces)
    if wsdl_ref_element is not None:
        wsdl_relative_path = wsdl_ref_element.attrib.get('ref', '')
        return wsdl_relative_path
    return ''
    
def extract_wsdl_operations(wsdl_path):
    operations = set()  # Utilizamos un conjunto en lugar de una lista
    if wsdl_path.endswith('.WSDL') and os.path.isfile(wsdl_path):
        with open(wsdl_path, 'r', encoding="utf-8") as f:
            wsdl_content = f.read()
            # Buscamos todas las coincidencias de "<operation name=" seguidas por el nombre de la operación
            operation_names = re.findall(r'operation name="([^"]+)', wsdl_content)
            for operation_name in operation_names:
                operations.add(operation_name)  # Agregamos el nombre de la operación al conjunto
    return list(operations)  # Convertimos el conjunto de vuelta a lista antes de devolverlo
  
def extract_osb_services_with_http_provider_id(project_path):

    osb_services = []
    elementos_xsd = []
    #st.success(f"project_path: {project_path}")
    for root, dirs, files in os.walk(project_path):
        if os.path.basename(root) == "Proxies":
            #st.success(f"✅ Proxies {elementos_xsd}")
            for file in files:
                if file.endswith('.ProxyService'):
                    osb_file_path = os.path.join(root, file)
                    #st.success(f"✅ osb_file_path {osb_file_path}")
                    project_name = extract_project_name_from_proxy(osb_file_path)
                    #st.success(f"✅ project_name {project_name}")
                    if project_name is None:
                        continue 
                    pipeline_path = extract_pipeline_path_from_proxy(osb_file_path, project_path)
                    #st.success(f"✅ pipeline_path {pipeline_path}")
                    with open(osb_file_path, 'r', encoding="utf-8") as f:
                        content = f.read()
                        if has_http_provider_id(content):
                            service_name = os.path.splitext(file)[0]
                            st.success(f"✅ service_name {service_name}")
                            service_url = extract_service_url(content)
                            st.success(f"✅ service_url {service_url}")
                            wsdl_relative_path = extract_wsdl_relative_path(content)
                            st.success(f"file: {file}")
                            st.success(f"project_path: {project_path}")
                            st.success(f"project_name: {project_name}")
                            st.success(f"pipeline_path: {pipeline_path}")
                            
                            #operacion_business = service_refs_ruta_pipeline(pipeline_path,project_path)
                            operacion_business = ""
                            st.success(f"operacion_business: {operacion_business}")
                            
                            st.success(f"service_name: {service_name}")
                            st.success(f"service_url: {service_url}")
                            st.success(f"wsdl_relative_path: {wsdl_relative_path}")
                            if wsdl_relative_path:
                                wsdl_path = os.path.join(project_path, wsdl_relative_path + ".WSDL")
                                capa_proyecto = '/'+ wsdl_relative_path.split('/')[0]
                                print_with_line_number("")
                                st.success(f"capa_proyecto: {capa_proyecto}")
                                print_with_line_number("")
                                st.success(f"wsdl_path: {wsdl_path}")
                                operations = extract_wsdl_operations(wsdl_path)
                                st.success(f"operations: {operations}")
                                imports = extract_xsd_import_paths(wsdl_path)
                                st.success(f"imports: {imports}")
                                #print_with_line_number("_________PRUEBA__________")
                                #print_with_line_number("")
                                
                                
                                # Crear un diccionario de mapeo entre operations y imports basado en nombres de archivos
                                operation_to_xsd = {}
                                for operation in operations:
                                    for xsd in imports:
                                        if operation.lower() in os.path.basename(xsd).lower():
                                            operation_to_xsd[operation] = xsd
                                            break
                                        else:
                                            xsd_names = [os.path.basename(xsd) for xsd in imports]  # Obtener solo los nombres de archivos XSD
                                            closest_match = difflib.get_close_matches(operation, xsd_names, n=1, cutoff=0.5)  # Buscar el más similar

                                            if closest_match:
                                                matched_xsd = next(xsd for xsd in imports if os.path.basename(xsd) == closest_match[0])
                                                operation_to_xsd[operation] = matched_xsd
                                            else:
                                                operation_to_xsd[operation] = None  # No se encontró una coincidencia
                                        
                                
                                st.success(f"operation_to_xsd: {operation_to_xsd}")
                                # Iterar sobre el diccionario y realizar la llamada a parse_xsd_file
                                for operation_name, xsd in operation_to_xsd.items():
                                    #print_with_line_number("")
                                    operation_actual = operation_name
                                    if operation_name == 'consultarProductoRequiereFirmaV21':
                                        st.success(f"operation_actual: {operation_actual}")
                                        st.success(f"service_name: {service_name}")
                                        st.success(f"operation_name: {operation_name}")
                                        st.success(f"xsd: {xsd}")
                                        st.success(f"service_url: {service_url}")
                                        st.success(f"capa_proyecto: {capa_proyecto}")
                                        st.success(f"operacion_business: {operacion_business}")
                                        xsd = os.path.splitext(xsd)[0] + ".XMLSchema"
                                        #print_with_line_number("")
                                        #print_with_line_number("")
                                    
                                        elementos_xsd = parse_xsd_file(project_path,xsd, operation_name,service_url,capa_proyecto,operacion_business,operations, service_name, operation_actual)
                                        st.success(f"elementos_xsd: {elementos_xsd}")
                                        #elementos_completos = list(elementos_xsd) + list(operations) + [operation_actual]
                                        osb_services.append(elementos_xsd)

    
    st.success(f"osb_services: {osb_services}")
    return osb_services

def extraer_jar(archivo_jar):
    """ Extrae el contenido de un .jar en una carpeta temporal en Windows. """
    try:
        # Obtener la ruta temporal
        ruta_temporal = os.path.join(tempfile.gettempdir(), "extraccion_jar")
        os.makedirs(ruta_temporal, exist_ok=True)

        # Verificar si el archivo JAR existe
        if not os.path.exists(archivo_jar):
            raise FileNotFoundError(f"El archivo .jar no existe: {archivo_jar}")

        # Ejecutar el comando 'jar xf'
        comando = f'java -jar {archivo_jar} -xf'
        resultado = subprocess.run(
            comando, shell=True, cwd=ruta_temporal,
            capture_output=True, text=True
        )

        # Validar la ejecución
        if resultado.returncode != 0:
            raise Exception(resultado.stderr)

        return ruta_temporal
    except Exception as e:
        st.error(f"Error al extraer el .jar: {e}")
        return None

def generar_documentacion(jar_path, plantilla_path):
    """Función que ejecuta la generación de documentación."""
    
    # Extraer ruta del proyecto desde el .jar
    jdeveloper_projects_dir = jar_path
    
    st.success(f"✅ jdeveloper_projects_dir {jdeveloper_projects_dir}")
    
    if not jdeveloper_projects_dir:
        st.error("No se pudo determinar la ruta del proyecto desde el .jar.")
        return
    
    # Cargar el documento de la plantilla
    doc = Document(plantilla_path)
    
    # Crear una carpeta temporal para almacenar los documentos
    temp_dir = tempfile.TemporaryDirectory()
    ruta_temporal = temp_dir.name  # Obtener la ruta temporal
    
    # Llamar a la función principal de tu script
    services_with_data = extract_osb_services_with_http_provider_id(jdeveloper_projects_dir)
    
    st.success(f"✅ services_with_data {services_with_data}")
    
    es_type = False
    
    # Initialize an empty set to store unique operation names
    operation_names = set()

    # Iterate through each tuple of request and response elements in services_with_data
    for request_elements, response_elements in services_with_data:
        # Iterate through each element in request_elements and response_elements
        for element in request_elements + response_elements:
            if 'Type' in element['elemento']:
                es_type = True
            #operation_name = element['elemento'].replace('Request', '').replace('Response', '').replace('Type', '')
            #st.success(f"operation_name: {operation_name}")
            service_name = element['service_name']
            # Agregar todas las operaciones de la lista 'operations'
            if 'operations' in element:
                operation_names.update(element['operations'])  # Agrega todas las operaciones a operation_names

    # Convert the set to a sorted list to get the operation names in alphabetical order
    unique_operations = sorted(operation_names)
    
    operaciones_formateadas = "\n".join(f"* {op}" for op in unique_operations)
    
    st.success(f"unique_operations: {unique_operations}")
    
    st.success(f"✅ unique_operations {unique_operations}")
    
    operation_elements = {}
    
    
    # Iterate through each unique operation
    for operation in unique_operations:
        if es_type:
            request_key = f"{operation}RequestType"
            response_key = f"{operation}ResponseType"
        else:
            request_key = f"{operation}Request"
            response_key = f"{operation}Response"
        
        # Initialize lists to store request and response elements for the current operation
        request_elements = []
        response_elements = []
        url_elements = []
        capa_proyecto = []
        business_elements = []
        
        # Iterate through services_with_data to find matching elements
        for request_data, response_data in services_with_data:
            # Check for request elements
            for element in request_data:
                if request_key in element['elemento']:
                    request_elements.append({'name': element['name'], 'type': element['type']})
                    url_elements.append({'url': element['url']})
                    capa_proyecto.append({'ruta': element['ruta']})
                    business_elements.append({'business': element['business']})
                    service_name = element['service_name']
            
            # Check for response elements
            for element in response_data:
                if response_key in element['elemento']:
                    response_elements.append({'name': element['name'], 'type': element['type']})
                    service_name = element['service_name']
        
        # Store the collected elements in the dictionary
        operation_elements[operation] = {
            'request': request_elements,
            'response': response_elements,
            'url': url_elements,
            'ruta': capa_proyecto, 
            'business': business_elements,
            'service_name': service_name
        }
        
    st.success(f"operation_elements: {operation_elements}")
    #st.success(f"service_name: {service_name}")
    # Print the result
    for operation, elements in operation_elements.items():
        
        if elements['request']:
            
            contiene_cabecera_entrada = False
            contiene_cabecera_salida = False
            
            if any('cabeceraEntrada.seguridad' in elem['name'] for elem in elements['request']):
                print_with_line_number("Se encontró al menos un elemento con '.cabeceraEntrada.seguridad'")
                contiene_cabecera_entrada = True
            
            if any('cabeceraSalida.' in elem['name'] for elem in elements['response']):
                contiene_cabecera_salida = True
            
            # Contar el número de tablas en el documento
            num_tables = len(doc.tables)
            
            st.success(f"El documento contiene {num_tables} tabla(s).")

            # Mostrar cada tabla
            for i, table in enumerate(doc.tables):
                st.success(f"\nTabla {i+1}:")
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    print_with_line_number('\t'.join(row_data))
            
            url = ""
            ruta =""
            business = ""
            
            for elem in elements['url']:
                url = elem['url']
                
            for elem in elements['ruta']:
                ruta = elem['ruta']
            
            for elem in elements['business']:
                business = elem['business']
                
            st.success(f"url: {url}")
            print_with_line_number("")
            st.success(f"ruta: {ruta}")
            print_with_line_number("")
            st.success(f"business: {business}")
            print_with_line_number("")
            fecha_actual = datetime.now()
            fecha_formateada = fecha_actual.strftime("%d/%m/%Y")
            
            print_with_line_number("")
            print_with_line_number("")
            st.success(f"operation: {operation}")
            
            st.success(f"elements: {elements}")
            print_with_line_number("")
            print_with_line_number("")
            
            # Definir las variables y sus valores
            variables = {
                '{nombre_servicio_inicial}': service_name,
                '{nombre_servicio_secundario}': service_name,
                '{nombre_servicio}': service_name,
                '{nombre_operacion_inicial}' : operation,
                '{nombre_operacion}': operation,
                '{unique_operations}': operaciones_formateadas,
                '{nombre_servicio_contrato}': service_name,
                '{nombre_servicio_wsdl}': service_name,
                '{nombre_servicio_contrato2}': service_name,
                '{nombre_servicio_tabla}': operation,
                '{fecha}': fecha_formateada,
                '{autor_inicial}': 'Kevin Torres',
                '{autor}': 'Kevin Torres',
                '{autor2}': 'Julian Orjuela',
                '{url}': url,
                '{operacion_legado}': business,
                '{proyecto_abc}': 'TENENCIA_COMPORTAMIENTO_ABC'
                # Añade más variables según sea necesario
            }
            
            st.success(f"service_name: {service_name}")
            st.success(f"variables: {variables}")
            
            tabla_cabecera_entrada_numero = 4
            tabla_cabecera_entrada = doc.tables[tabla_cabecera_entrada_numero - 1]  # Las tablas se indexan desde 0, por eso restamos 1

            tabla_request_numero = 5
            tabla_request = doc.tables[tabla_request_numero - 1]  # Las tablas se indexan desde 0, por eso restamos 1
            
            tabla_cabecera_salida_numero = 6
            tabla_cabecera_salida = doc.tables[tabla_cabecera_salida_numero - 1]  # Las tablas se indexan desde 0, por eso restamos 1
            
            # Listas para almacenar las filas de cada subtabla
            cabecera_salida = []
            datos_respuesta = []
            
            # Variables de control
            seccion_actual = None
            
            st.success(f"Número total de tablas en el documento: {len(doc.tables)}")
            
            for i, table in enumerate(doc.tables):
                st.success(f"Tabla {i + 1}:")  # Mostrar el número de la tabla

                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]  # Extraer el texto de cada celda
                    st.success(f"  {row_text}")  # Imprimir el contenido de la fila

                print_with_line_number("-" * 50)  # Separador entre tablas
           
           
            # Recorrer las filas de la tabla 7
            for row in tabla_cabecera_salida.rows:
                row_text = [cell.text.strip() for cell in row.cells]

                # Detectar la cabecera de cada subtabla
                if "Cabecera de salida" in row_text:
                    seccion_actual = "cabecera_salida"
                    continue  # Saltar a la siguiente fila

                if "Datos Respuesta" in row_text:
                    seccion_actual = "datos_respuesta"
                    continue  # Saltar a la siguiente fila

                # Guardar las filas en la subtabla correspondiente
                if seccion_actual == "cabecera_salida":
                    cabecera_salida.append(row_text)

                elif seccion_actual == "datos_respuesta":
                    datos_respuesta.append(row_text)
           
            # Identificar la sección "Datos Respuesta"
            for row in tabla_cabecera_salida.rows:
                if "Datos Respuesta" in row.cells[0].text:
                    tabla_response = tabla_cabecera_salida  # Ahora sí es una tabla válida
                    break
            else:
                print_with_line_number("No se encontró la sección 'Datos Respuesta' en la tabla 7.")
                tabla_response = None  # Para evitar futuros errores
           
            
            # Datos por defecto para LONGITUD y OBSERVACIÓN
            default_longitud = "default"
            default_observacion = ""
            
            # Limpiar la tabla antes de agregar elementos de esta operación
            if not contiene_cabecera_entrada:
                tbl = tabla_cabecera_entrada._element
                tbl.getparent().remove(tbl)
                while len(tabla_cabecera_entrada.rows) > 1:
                    tabla_cabecera_entrada._element.remove(tabla_cabecera_entrada.rows[1]._element)
                    
            # Limpiar la tabla antes de agregar elementos de esta operación
            if not contiene_cabecera_salida:
                #tbl = tabla_cabecera_salida._element
                #tbl.getparent().remove(tbl)
                while len(tabla_cabecera_salida.rows) > 1:
                    tabla_cabecera_salida._element.remove(tabla_cabecera_salida.rows[1]._element)

            # Limpiar la tabla antes de agregar elementos de esta operación
            while len(tabla_request.rows) > 2:
                tabla_request._element.remove(tabla_request.rows[2]._element)
            
            # Procesar los datos
            for elem in elements['request']:
                
                if 'cabeceraEntrada.' not in elem['name']:
                    # Añadir una nueva fila al final de la tabla
                    fila = tabla_request.add_row().cells
                    
                    # Rellenar la fila con los datos correspondientes
                    #fila[0].text = operation + "Request" + "." + elem['name']
                    fila[0].text = elem['name']
                    st.success(f"fila[0].text: {fila[0].text}")
                    fila[1].text = elem['name']
                    campo = fila[1].text.split('.')[-1]
                    fila[1].text = campo
                    st.success(f"fila[1].text: {fila[1].text}")
                    fila[2].text = default_longitud
                    fila[3].text = elem['type']
                    tipo_campo = fila[3].text.split(':')[-1]
                    if tipo_campo == 'string':
                        tipo_campo = 'Alfanumérico'
                    fila[3].text = tipo_campo
                    st.success(f"fila[3].text: {fila[3].text}")
                
                
            # Limpiar la tabla antes de agregar elementos de esta operación
            while len(tabla_response.rows) > 2:
                tabla_response._element.remove(tabla_response.rows[2]._element)
            
            # Procesar los datos
            for elem in elements['response']:
                
                
                if 'cabeceraSalida.' not in elem['name']:
                    # Añadir una nueva fila al final de la tabla
                    fila = tabla_response.add_row().cells
                    
                    # Rellenar la fila con los datos correspondientes
                    #fila[0].text = operation + "Response" + "." + elem['name']
                    fila[0].text = elem['name']
                    st.success(f"fila[0].text: {fila[0].text}")
                    fila[1].text = elem['name']
                    campo = fila[1].text.split('.')[-1]
                    fila[1].text = campo
                    st.success(f"fila[1].text: {fila[1].text}")
                    fila[2].text = default_longitud
                    fila[3].text = elem['type']
                    tipo_campo = fila[3].text.split(':')[-1]
                    if tipo_campo == 'string':
                        tipo_campo = 'Alfanumérico'
                    fila[3].text = tipo_campo
                    st.success(f"fila[3].text: {fila[3].text}")

            print_with_line_number("___________________________________________")
            
            st.success(f"✅ temp_dir  {temp_dir }")
            st.success(f"✅ ruta_temporal  {ruta_temporal }")

            # Lista para almacenar las rutas de los documentos generados
            documentos_generados = []

            ruta_proyecto = ruta.strip("/")  # Asegurar que la ruta no tenga "/" al inicio
            nombre_documento = f"Especificación Servicio WSDL {operation}.docx"
            
            # Crear la ruta dentro de la carpeta temporal
            carpeta_destino = os.path.join(ruta_temporal, ruta_proyecto)
            os.makedirs(carpeta_destino, exist_ok=True)  # Crear la carpeta si no existe
            
            ruta_guardado = os.path.join(carpeta_destino, nombre_documento)
            doc_nuevo = replace_text_in_doc(doc, variables)
            doc_nuevo.save(ruta_guardado)  # Guardar en la carpeta temporal
                
            # Agregar a la lista de documentos generados
            documentos_generados.append((ruta_guardado, os.path.join(ruta_proyecto, nombre_documento)))
            st.success(f"Documento guardado: {ruta_guardado}")

            # Crear el archivo ZIP en memoria
            zip_buffer = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
            zip_path = zip_buffer.name  # Obtener la ruta del archivo ZIP

            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
                for doc_path, zip_relative_path in documentos_generados:
                    zipf.write(doc_path, zip_relative_path)  # Mantener la estructura dentro del ZIP

            st.success(f"ZIP creado: {zip_path}")

            # Permitir la descarga del ZIP desde Streamlit
            with open(zip_path, "rb") as file:
                zip_bytes = file.read()

            st.download_button(
                label="📥 Descargar TODOS los documentos en ZIP",
                data=zip_bytes,
                file_name="Documentos_OSB.zip",
                mime="application/zip"
            )
                
    st.success("Documentación generada con éxito!")

def main():
    st.title("📄 Generador de Documentación OSB")
    
    jar_file = st.file_uploader("Sube el archivo .jar con dependencias", type=["jar"])
    plantilla_file = st.file_uploader("Sube la plantilla de Word", type=["docx"])
    #destino_path = st.text_input("Ruta donde se generarán los documentos")
    
    if jar_file:
        jar_path = "temp.jar"

        # Guardar el archivo
        with open(jar_path, "wb") as f:
            f.write(jar_file.getbuffer())

        # Ruta donde se extraerán los archivos
        carpeta_destino = "extraccion_jar"

        # Extraer los archivos del .jar
        try:
            with zipfile.ZipFile(jar_path, "r") as jar:
                jar.extractall(carpeta_destino)
                archivos_extraidos = jar.namelist()
            
            st.success(f"✅ Archivos extraídos en: {carpeta_destino}")
            st.write("📂 Archivos extraídos:")
            st.write(archivos_extraidos)
        except zipfile.BadZipFile:
            st.error("❌ Error: El archivo no es un JAR válido o está dañado.")
    
    if st.button("Generar Documentación"):
        if jar_file and plantilla_file:
            with st.spinner("Generando documentación..."):
                generar_documentacion(carpeta_destino, plantilla_file)
        else:
            st.error("Por favor, sube todos los archivos y proporciona la ruta de destino.")

if __name__ == "__main__":
    main()
